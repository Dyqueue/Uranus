from urllib.parse import urljoin
from bs4 import BeautifulSoup
import urllib.request as req
import os.path
from selenium import webdriver
from urllib.request import urlopen
from docx import Document
from docx.shared import Inches
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
document = Document()
document.add_heading(u'Python 操作Word実例', 3)
p_total = document.add_heading()
n=[15,16,17,18,19,20]
num = 0
for i in n:
    Url="http://www.city.kumamoto.jp/loc/pub/default.aspx?c_id={0}".format(i)
    browser=webdriver.Chrome()
    browser.implicitly_wait(3)
    res=req.urlopen(Url)
    soup=BeautifulSoup(res,"html.parser")
    for i in soup.select(" td.name "):
        h=i.select("a")
        for z in h:
            href=z.attrs["href"]
            urla=urljoin(Url,href)
            browser.get(urla)
            text=browser.find_element_by_id("mainBaseInfo")
            html=urlopen(urla)
            bsObj = BeautifulSoup(html, 'html.parser')
            t1 = bsObj.find_all('a')
            basket=[]
            for t2 in t1:
                t3 = t2.get('href')
                basket.append(t3)
            t4=basket[-11]
            print(z.text)
            document.add_heading(z.text, 2)
            print(text.text)
            document.add_paragraph(text.text)
            print(t4)
            document.add_paragraph(t4)
            browser.get(t4)
            browser.implicitly_wait(10)
            num += 1
            browser.save_screenshot("web{0}.png".format(num))
            document.add_picture("web{0}.png".format(num), width=Inches(2.5))
            document.save("C:/Users/iwaao/Anaconda3/Lib/site-packages/Test.docx")
if __name__ == '__main__':
    fromaddr = 'iwaao@yahoo.com'
    password = 'dyq950217'
    toaddrs = ['iwaao@yahoo.com', 'iwaao@yahoo.com']
    message = MIMEText('testプログラムです。','plain','utf-8')
    docxFile = 'C:/Users/iwaao/Anaconda3/Lib/site-packages/Test.docx'
    docxApart = MIMEApplication(open(docxFile, 'rb').read())
    docxApart.add_header('Content-Disposition', 'attachment', filename=docxFile)
    m = MIMEMultipart()
    m.attach(docxApart)
    m['Subject'] = 'title'
    try:
        server = smtplib.SMTP('pop.mail.yahoo.com')
        server.login(fromaddr,password)
        server.sendmail(fromaddr, toaddrs, m.as_string())
        print('success')
